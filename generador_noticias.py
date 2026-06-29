"""
Generador de Noticias — MJR-21
Lee un Word (.docx) o PDF con negrillas/cursivas/párrafos/tablas/imágenes y genera
la carpeta nuevos/{slug}/ lista, luego corre build.py.
Pestañas: Crear noticia | Editar noticia | Eliminar / deshacer
"""

import platform
import re
import shutil
import subprocess
import threading
import unicodedata
import zipfile
from datetime import date, datetime
from pathlib import Path
from tkinter import filedialog, messagebox

import customtkinter as ctk
from docx import Document
from docx.oxml.ns import qn
import fitz  # pymupdf

ROOT = Path(__file__).parent
NUEVOS = ROOT / "nuevos"
OUTPUT_ARTICULOS = ROOT / "articulos"
IMG_ARTICULOS = ROOT / "img" / "articulos"

ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")


# ══════════════════════════════════════════════════════════════════════════════
#  UTILIDADES GENERALES
# ══════════════════════════════════════════════════════════════════════════════

def slugify(texto):
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode()
    texto = texto.lower()
    texto = re.sub(r"[^a-z0-9\s-]", "", texto)
    texto = re.sub(r"\s+", "-", texto.strip())
    return texto[:60].strip("-")


def validar_fecha(texto):
    try:
        datetime.strptime(texto.strip(), "%Y-%m-%d")
        return True
    except ValueError:
        return False


def elegir_archivo(titulo, filtro_zenity, filtros_tk):
    if platform.system() == "Linux" and shutil.which("zenity"):
        try:
            resultado = subprocess.run(
                ["zenity", "--file-selection", "--title", titulo, "--file-filter", filtro_zenity],
                capture_output=True, text=True,
            )
            ruta = resultado.stdout.strip()
            return ruta if ruta else None
        except Exception:
            pass
    return filedialog.askopenfilename(title=titulo, filetypes=filtros_tk) or None


# ══════════════════════════════════════════════════════════════════════════════
#  CONVERSIÓN DOCX → HTML
# ══════════════════════════════════════════════════════════════════════════════

def runs_a_html(runs):
    partes = []
    for run in runs:
        texto = run.text
        if not texto:
            continue
        t = texto.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if run.bold:
            t = f"<strong>{t}</strong>"
        if run.italic:
            t = f"<em>{t}</em>"
        if run.underline:
            t = f"<u>{t}</u>"
        partes.append(t)
    return "".join(partes).strip()


def parrafo_a_html(parrafo):
    return runs_a_html(parrafo.runs)


def tabla_a_filas(tabla, filas_pendientes):
    for i, fila in enumerate(tabla.rows):
        celdas = []
        for celda in fila.cells:
            contenido = " ".join(
                runs_a_html(p.runs) for p in celda.paragraphs if p.text.strip()
            )
            es_header = (i == 0 and not filas_pendientes)
            tc = "th" if es_header else "td"
            celdas.append(f"<{tc}>{contenido}</{tc}>")
        filas_pendientes.append(f"<tr>{''.join(celdas)}</tr>")


def extraer_imagenes_docx(ruta_docx, carpeta_destino):
    carpeta_img = carpeta_destino / "img_inline"
    carpeta_img.mkdir(parents=True, exist_ok=True)
    doc = Document(str(ruta_docx))
    rels = doc.part.rels
    mapa = {}
    with zipfile.ZipFile(str(ruta_docx)) as z:
        for rid, rel in rels.items():
            if "image" not in rel.reltype:
                continue
            zip_path = f"word/{rel.target_ref}"
            if zip_path not in z.namelist():
                continue
            ext = Path(rel.target_ref).suffix.lower() or ".png"
            nombre = f"{rid}{ext}"
            destino = carpeta_img / nombre
            with z.open(zip_path) as src, open(destino, "wb") as dst:
                dst.write(src.read())
            mapa[rid] = destino
    return mapa


def imagen_inline_html(parrafo_elem, mapa_rid, slug):
    blips = parrafo_elem.findall('.//' + qn('a:blip'))
    if not blips:
        return None
    rid = blips[0].get(qn('r:embed'))
    if not rid or rid not in mapa_rid:
        return None
    ruta_img = mapa_rid[rid]
    ruta_publica = f"/img/articulos/{slug}/img_inline/{ruta_img.name}"
    return f'<figure class="article-inline-img"><img src="{ruta_publica}" alt="Imagen del artículo" loading="lazy"></figure>'


def docx_a_cuerpo(ruta_docx, carpeta_slug, slug):
    doc = Document(str(ruta_docx))
    parrafo_map = {p._element: p for p in doc.paragraphs}
    tabla_map   = {t._element: t for t in doc.tables}
    mapa_rid = extraer_imagenes_docx(ruta_docx, carpeta_slug) if carpeta_slug else {}

    bloques = []
    primer_elemento = True
    filas_pendientes = []

    def volcar_tabla():
        if filas_pendientes:
            bloques.append(
                '<div class="table-wrap"><table>' + "".join(filas_pendientes) + '</table></div>'
            )
            filas_pendientes.clear()

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            img_html = imagen_inline_html(child, mapa_rid, slug)
            if img_html:
                volcar_tabla()
                bloques.append(img_html)
                primer_elemento = False
                continue

            parrafo = parrafo_map.get(child)
            if parrafo is None or not parrafo.text.strip():
                primer_elemento = False
                continue

            if primer_elemento:
                primer_elemento = False
                if parrafo.style.name.startswith("Heading") or parrafo.style.name == "Title":
                    continue

            volcar_tabla()
            html = parrafo_a_html(parrafo)
            if html:
                estilo = parrafo.style.name
                if estilo.startswith("Heading 1"):
                    bloques.append(f"<h2>{html}</h2>")
                elif estilo.startswith("Heading"):
                    bloques.append(f"<h3>{html}</h3>")
                else:
                    bloques.append(f"<p>{html}</p>")
            primer_elemento = False

        elif tag == "tbl":
            tabla = tabla_map.get(child)
            if tabla is not None:
                tabla_a_filas(tabla, filas_pendientes)
            primer_elemento = False
        else:
            primer_elemento = False

    volcar_tabla()
    return "\n\n".join(bloques)


def primer_titulo(ruta_docx):
    doc = Document(str(ruta_docx))
    for p in doc.paragraphs:
        if p.text.strip():
            if p.style.name.startswith("Heading") or p.style.name == "Title":
                return p.text.strip()
            break
    return ""


# ══════════════════════════════════════════════════════════════════════════════
#  CONVERSIÓN PDF → HTML
# ══════════════════════════════════════════════════════════════════════════════

def _rgb(color_int):
    return ((color_int >> 16) & 0xFF, (color_int >> 8) & 0xFF, color_int & 0xFF)


def _cerca(c1, c2, tol=25):
    return all(abs(a - b) <= tol for a, b in zip(c1, c2))


# Colores exactos del PDF MJR-21 (verificados span a span)
_AZUL    = (26,  63,  160)   # encabezados I/IV/VII/VIII + col izquierda tabla
_VERDE   = (26,  122, 60)    # encabezados II/V/VIII
_AMARILLO= (245, 197, 24)    # encabezado X
_ROJO    = (192, 57,  43)    # encabezados III/VI/IX + subsecciones
_AZUL2   = (46,  74,  158)   # subsecciones alternativas (2.1)
_BLANCO  = (255, 255, 255)   # texto en header de tabla (fondo azul)
_OSCURO  = (30,  30,  46)    # texto de cuerpo (col derecha tabla + párrafos)
_GRIS    = (136, 136, 136)   # pie de página

# X fijo de separación de columnas (verificado en el PDF real)
# Col izquierda: x ≈ 72  |  Col derecha: x ≈ 202
_COL_SPLIT = 190  # cualquier span con x >= 190 es columna derecha


def _es_h_romano(rgb, sz):
    """Encabezados tipo I., II., ..., X. — tamaño 15, cualquier color de sección."""
    if sz < 14:
        return False
    return (_cerca(rgb, _AZUL) or _cerca(rgb, _VERDE) or
            _cerca(rgb, _ROJO) or _cerca(rgb, _AMARILLO, 30))


def _es_subseccion(rgb, sz):
    """Subsecciones tipo 1.1, 2.2 — tamaño ~12, rojo o azul2."""
    if sz < 11:
        return False
    return _cerca(rgb, _ROJO) or _cerca(rgb, _AZUL2, 20)


def _es_col_izq(rgb, bold, x):
    """Celda izquierda de tabla: azul bold, x < _COL_SPLIT."""
    return x < _COL_SPLIT and bold and _cerca(rgb, _AZUL)


def _es_col_der(rgb, x):
    """Celda derecha de tabla: oscuro, x >= _COL_SPLIT."""
    return x >= _COL_SPLIT and (_cerca(rgb, _OSCURO, 30) or
                                 (rgb[0] < 80 and rgb[1] < 80 and rgb[2] < 80))


def _es_gris(rgb):
    return _cerca(rgb, _GRIS, 30)


def _escape(t):
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _es_logo_img(bbox, page_h, iw, ih):
    x0, y0, x1, y1 = bbox
    en_margen = y0 < page_h * 0.22 or y1 > page_h * 0.82
    return en_margen and (iw < 200 or ih < 200)


# Patrones de texto para detección de estructura
_PAT_ROMANO    = re.compile(r'^[IVXLCDM]+\.\s+\S', re.IGNORECASE)
_PAT_SUBSEC    = re.compile(r'^\d+\.\d+[\s\.]')
_PAT_PAG       = re.compile(r'^\s*(Pág\.\s*\d+|MJR-21\s*•|\d+\s*$)')
_PAT_HDR_TABLA = re.compile(r'DISPOSICI[OÓ]N|CONTENIDO\s+DE\s+LA', re.IGNORECASE)


def pdf_a_cuerpo(ruta_pdf, carpeta_slug, slug):
    """
    Convierte PDF MJR-21 → HTML.
    La portada (pág 0) se salta completamente.
    Las tablas se reconstruyen fila por fila usando la X exacta de cada span.
    El estado (tabla abierta, encabezados parciales) persiste entre páginas.
    """
    doc = fitz.open(str(ruta_pdf))
    output  = []
    xrefs_vistos = set()
    carpeta_img  = None
    img_counter  = 0

    if carpeta_slug:
        carpeta_img = carpeta_slug / "img_inline"
        carpeta_img.mkdir(parents=True, exist_ok=True)

    # Estado persistente entre páginas
    en_tabla    = False
    filas_tabla = []      # filas HTML acumuladas de la tabla actual
    buf_izq     = []      # líneas de texto de la celda izquierda actual
    buf_der     = []      # líneas de texto de la celda derecha actual
    buf_h2      = []      # palabras del encabezado romano en curso
    buf_h3      = []      # palabras de la subsección en curso

    def flush_h2():
        if buf_h2:
            txt = " ".join(buf_h2).strip()
            output.append(
                f'<h2 style="color:#1a3fa0;border-left:5px solid #1a3fa0;'
                f'padding-left:0.7em;margin:1.8em 0 0.5em;font-size:1.25rem;font-weight:700">'
                f'{_escape(txt)}</h2>'
            )
            buf_h2.clear()

    def flush_h3():
        if buf_h3:
            txt = " ".join(buf_h3).strip()
            output.append(
                f'<h3 style="color:#c0392b;margin:1.2em 0 0.3em;font-size:1.05rem;font-weight:700">'
                f'{_escape(txt)}</h3>'
            )
            buf_h3.clear()

    def flush_fila():
        """Cierra la fila actual de tabla y la agrega a filas_tabla."""
        if buf_izq or buf_der:
            izq = " ".join(buf_izq).strip()
            der = " ".join(buf_der).strip()
            filas_tabla.append(
                f"<tr>"
                f'<td><strong style="color:#1a3fa0">{_escape(izq)}</strong></td>'
                f"<td>{_escape(der)}</td>"
                f"</tr>"
            )
            buf_izq.clear()
            buf_der.clear()

    def flush_tabla():
        nonlocal en_tabla
        flush_fila()
        if filas_tabla:
            hdr = (
                "<tr>"
                '<th style="background:#1a3fa0;color:#fff;padding:.5rem .9rem;text-align:left">DISPOSICIÓN CONSTITUCIONAL</th>'
                '<th style="background:#1a3fa0;color:#fff;padding:.5rem .9rem;text-align:left">CONTENIDO DE LA PROPUESTA</th>'
                "</tr>"
            )
            output.append(
                f'<div class="table-wrap"><table>{hdr}{"".join(filas_tabla)}</table></div>'
            )
            filas_tabla.clear()
        en_tabla = False

    # ── Iterar páginas ─────────────────────────────────────────────────────

    for num_pag, pagina in enumerate(doc):
        page_h = pagina.rect.height
        page_w = pagina.rect.width

        # SALTAR PORTADA COMPLETA (pág 0)
        if num_pag == 0:
            continue

        # ── Imágenes ───────────────────────────────────────────────────────
        if carpeta_img:
            for img_info in pagina.get_images(full=True):
                xref = img_info[0]
                if xref in xrefs_vistos:
                    continue
                base = doc.extract_image(xref)
                iw, ih = base["width"], base["height"]
                bboxes = pagina.get_image_rects(xref)
                bbox = bboxes[0] if bboxes else None
                if bbox and _es_logo_img(bbox, page_h, iw, ih):
                    xrefs_vistos.add(xref); continue
                if iw < 80 or ih < 80:
                    xrefs_vistos.add(xref); continue
                xrefs_vistos.add(xref)
                nombre = f"pdf_img_{img_counter}.{base['ext']}"
                (carpeta_img / nombre).write_bytes(base["image"])
                ruta_pub = f"/img/articulos/{slug}/img_inline/{nombre}"
                output.append(
                    f'<figure class="article-inline-img">'
                    f'<img src="{ruta_pub}" alt="Imagen" loading="lazy"></figure>'
                )
                img_counter += 1

        # ── Extraer spans de la página ─────────────────────────────────────
        # Cada span: (y, x, sz, bold, rgb, text)
        raw_spans = []
        for blk in pagina.get_text("rawdict", flags=0)["blocks"]:
            if blk["type"] != 0:
                continue
            for linea in blk["lines"]:
                ly0 = linea["bbox"][1]
                for span in linea["spans"]:
                    chars = span.get("chars", [])
                    text  = "".join(c["c"] for c in chars).strip()
                    if not text:
                        continue
                    rgb  = _rgb(span["color"])
                    sz   = span["size"]
                    bold = bool(span["flags"] & 16)
                    x0   = span["bbox"][0]
                    raw_spans.append((ly0, x0, sz, bold, rgb, text))

        if not raw_spans:
            continue

        # Ordenar por Y luego X
        raw_spans.sort(key=lambda s: (round(s[0] / 4) * 4, s[1]))

        # Agrupar en "líneas lógicas" por Y±4px
        lineas_logicas = []
        for sp in raw_spans:
            y = sp[0]
            if lineas_logicas and abs(y - lineas_logicas[-1][0][0]) <= 4:
                lineas_logicas[-1].append(sp)
            else:
                lineas_logicas.append([sp])

        # ── Procesar cada línea lógica ─────────────────────────────────────
        for linea in lineas_logicas:
            # Propiedades del primer span de la línea
            y0, x0_primero, sz0, bold0, rgb0, _ = linea[0]

            # Texto completo de la línea
            texto_full = " ".join(sp[5] for sp in linea).strip()
            if not texto_full:
                continue

            # ── 1. Filtrar ruido ──────────────────────────────────────────
            if _es_gris(rgb0):
                continue
            if _PAT_PAG.match(texto_full):
                continue
            if len(texto_full) < 3 and not any(c.isalpha() for c in texto_full):
                continue
            # Encabezado del movimiento en zona header (azul pequeño, y < 100)
            if y0 < 100 and sz0 <= 10:
                continue

            # ── 2. Header de tabla (texto blanco) ────────────────────────
            if all(_cerca(sp[4], _BLANCO, 15) for sp in linea):
                # Es el encabezado azul de la tabla con texto blanco
                # Si había h2/h3 pendiente, flushearlo antes
                flush_h2()
                flush_h3()
                # Si ya había una tabla abierta, cerrarla (nueva tabla)
                if en_tabla:
                    flush_tabla()
                en_tabla = True
                continue

            # ── 3. Encabezado romano (I., II., ...) ──────────────────────
            # Detectar por PATRÓN de texto (independiente del color)
            if _PAT_ROMANO.match(texto_full.strip()) and sz0 >= 13:
                if en_tabla:
                    flush_tabla()
                flush_h2()
                flush_h3()
                # Iniciar nuevo h2
                buf_h2.append(texto_full)
                continue

            # ── 4. Continuación de encabezado romano partido ──────────────
            # (mismo color de sección, sz >= 13, sin ser inicio de nueva sección)
            if buf_h2 and not en_tabla and sz0 >= 13 and _es_h_romano(rgb0, sz0):
                buf_h2.append(texto_full)
                continue

            # ── 5. Subsección N.N Título ──────────────────────────────────
            if _PAT_SUBSEC.match(texto_full.strip()) and sz0 >= 11:
                flush_h3()
                flush_h2()
                buf_h3.append(texto_full)
                continue

            # ── 6. Continuación de subsección partida ─────────────────────
            if buf_h3 and not en_tabla and sz0 >= 11 and _es_subseccion(rgb0, sz0):
                buf_h3.append(texto_full)
                continue

            # ── 7. Contenido de tabla ─────────────────────────────────────
            if en_tabla:
                # Separar los spans de la línea en col izq y col der
                # usando la X EXACTA de cada span individual
                spans_izq = [sp for sp in linea if sp[1] < _COL_SPLIT]
                spans_der = [sp for sp in linea if sp[1] >= _COL_SPLIT]

                texto_izq = " ".join(sp[5] for sp in spans_izq).strip()
                texto_der = " ".join(sp[5] for sp in spans_der).strip()

                if texto_izq:
                    # Nueva celda izquierda → cerrar fila anterior
                    if buf_izq or buf_der:
                        flush_fila()
                    buf_izq.append(texto_izq)

                if texto_der:
                    buf_der.append(texto_der)

                continue

            # ── 8. Párrafo normal ─────────────────────────────────────────
            flush_h2()
            flush_h3()

            partes_html = []
            for sp in linea:
                t = _escape(sp[5])
                if sp[3]:   # bold
                    t = f"<strong>{t}</strong>"
                partes_html.append(t)
            output.append(f"<p>{''.join(partes_html)}</p>")

    # ── Cerrar todo al terminar el documento ──────────────────────────────
    if en_tabla:
        flush_tabla()
    flush_h2()
    flush_h3()

    doc.close()

    # ── Post-proceso: fusionar <p> consecutivos ────────────────────────────
    resultado = []
    buf_p = []

    def volcar_p():
        if buf_p:
            resultado.append("<p>" + " ".join(buf_p) + "</p>")
            buf_p.clear()

    for bloque in output:
        if bloque.startswith("<p>") and "<figure" not in bloque:
            inner = re.sub(r'^<p>(.*)</p>$', r'\1', bloque, flags=re.DOTALL)
            buf_p.append(inner)
        else:
            volcar_p()
            resultado.append(bloque)
    volcar_p()

    return "\n\n".join(resultado)


def primer_titulo_pdf(ruta_pdf):
    doc = fitz.open(str(ruta_pdf))
    candidatos = []
    for pagina in doc:
        for b in pagina.get_text("rawdict", flags=0)["blocks"]:
            if b["type"] != 0:
                continue
            for linea in b["lines"]:
                for span in linea["spans"]:
                    chars = span.get("chars", [])
                    text  = "".join(c["c"] for c in chars).strip()
                    if text and span["size"] >= 12:
                        rgb = _rgb(span["color"])
                        if not _es_gris(rgb) and not _cerca(rgb, _BLANCO, 15):
                            candidatos.append((span["size"], text))
        if candidatos:
            break
    doc.close()
    if candidatos:
        candidatos.sort(key=lambda x: -x[0])
        return candidatos[0][1]
    return ""


# ══════════════════════════════════════════════════════════════════════════════
#  UTILIDADES DE ARTÍCULOS
# ══════════════════════════════════════════════════════════════════════════════

def listar_articulos_existentes():
    resultado = []
    if not NUEVOS.exists():
        return resultado
    for carpeta in sorted(NUEVOS.iterdir()):
        if not carpeta.is_dir() or carpeta.name.startswith("_"):
            continue
        archivo_txt = carpeta / "articulo.txt"
        titulo = carpeta.name
        if archivo_txt.exists():
            for linea in archivo_txt.read_text(encoding="utf-8").splitlines():
                if linea.upper().startswith("TITULO:"):
                    titulo = linea.split(":", 1)[1].strip() or carpeta.name
                    break
        resultado.append((carpeta.name, titulo))
    return resultado


def leer_metadata_articulo(slug):
    archivo = NUEVOS / slug / "articulo.txt"
    if not archivo.exists():
        return {}
    texto = archivo.read_text(encoding="utf-8")
    metadata_texto, _, _ = texto.partition("---")
    meta = {}
    for linea in metadata_texto.strip().splitlines():
        if ":" in linea:
            clave, _, valor = linea.partition(":")
            meta[clave.strip().lower()] = valor.strip()
    return meta


def eliminar_articulo(slug):
    for carpeta in [NUEVOS / slug, OUTPUT_ARTICULOS / slug, IMG_ARTICULOS / slug]:
        if carpeta.exists():
            shutil.rmtree(carpeta)


def _es_pdf(ruta):
    return ruta is not None and Path(ruta).suffix.lower() == ".pdf"


def _convertir_a_cuerpo(ruta, carpeta_slug, slug):
    if _es_pdf(ruta):
        return pdf_a_cuerpo(ruta, carpeta_slug, slug)
    return docx_a_cuerpo(ruta, carpeta_slug, slug)


def _obtener_titulo_sugerido(ruta):
    if _es_pdf(ruta):
        return primer_titulo_pdf(ruta)
    return primer_titulo(ruta)


FILTRO_DOC_ZENITY = "Documentos | *.docx *.pdf"
FILTRO_DOC_TK     = [("Word o PDF", "*.docx *.pdf"), ("Word", "*.docx"), ("PDF", "*.pdf")]
FILTRO_IMG_ZENITY = "Imágenes | *.jpg *.jpeg *.png *.webp"
FILTRO_IMG_TK     = [("Imágenes", "*.jpg *.jpeg *.png *.webp")]


# ══════════════════════════════════════════════════════════════════════════════
#  VENTANA DE PREVISUALIZACIÓN
# ══════════════════════════════════════════════════════════════════════════════

class VentanaPreview(ctk.CTkToplevel):
    def __init__(self, master, titulo, categoria, autor, fecha, resumen, cuerpo_html,
                 carpeta_preview=None):
        super().__init__(master)
        self.title("Previsualización del artículo")
        self.geometry("820x720")
        self.resizable(True, True)
        self.resultado = False
        self._imagenes_tk = []
        self._carpeta_preview = carpeta_preview

        header = ctk.CTkFrame(self, fg_color="#0C4988", corner_radius=0)
        header.pack(fill="x")
        ctk.CTkLabel(
            header, text="👁  Previsualización — revisa antes de publicar",
            font=("Arial", 13, "bold"), text_color="white"
        ).pack(side="left", padx=16, pady=10)

        self._construir_vista(titulo, categoria, autor, fecha, resumen, cuerpo_html)

        bar = ctk.CTkFrame(self, fg_color="transparent")
        bar.pack(fill="x", padx=20, pady=12)
        ctk.CTkButton(
            bar, text="✅ Confirmar y publicar", command=self._confirmar,
            fg_color="#0B7439", hover_color="#075226", font=("Arial", 13, "bold"), height=40
        ).pack(side="left", expand=True, fill="x", padx=(0, 8))
        ctk.CTkButton(
            bar, text="✏️ Volver a editar", command=self._cancelar,
            fg_color="#8E0209", hover_color="#5e0106", font=("Arial", 13), height=40
        ).pack(side="left", expand=True, fill="x")

        self.update_idletasks()
        self.after(0, self.grab_set)
        self.protocol("WM_DELETE_WINDOW", self._cancelar)

    def _construir_vista(self, titulo, categoria, autor, fecha, resumen, cuerpo_html):
        import tkinter as tk

        frame = ctk.CTkFrame(self, fg_color="#f4efe2", corner_radius=0)
        frame.pack(fill="both", expand=True)

        sb = tk.Scrollbar(frame)
        sb.pack(side="right", fill="y")

        self._txt = tk.Text(
            frame, wrap="word", yscrollcommand=sb.set,
            bg="#f4efe2", fg="#2b2620", font=("Georgia", 13),
            relief="flat", padx=28, pady=20, cursor="arrow", state="normal"
        )
        self._txt.pack(fill="both", expand=True)
        sb.config(command=self._txt.yview)

        txt = self._txt
        txt.tag_config("tag_pill",  foreground="#8e0209", font=("Arial", 10), spacing1=4)
        txt.tag_config("h1",        font=("Arial", 22, "bold"), spacing1=10, spacing3=4)
        txt.tag_config("h2",        font=("Arial", 15, "bold"), spacing1=14, spacing3=3, foreground="#1a3fa0")
        txt.tag_config("h3",        font=("Arial", 13, "bold"), spacing1=10, spacing3=2, foreground="#c0392b")
        txt.tag_config("meta",      foreground="#888888", font=("Arial", 10), spacing3=10)
        txt.tag_config("p",         font=("Georgia", 13), spacing3=8)
        txt.tag_config("resumen",   font=("Georgia", 12, "italic"), foreground="#555555", spacing3=12)
        txt.tag_config("sep",       font=("Arial", 6), spacing1=6, spacing3=6)
        txt.tag_config("tabla_hdr", font=("Arial", 10, "bold"), background="#1a3fa0",
                       foreground="white", spacing1=4, spacing3=4)
        txt.tag_config("tabla_izq", font=("Arial", 11, "bold"), foreground="#1a3fa0", spacing1=2)
        txt.tag_config("tabla_der", font=("Georgia", 11), spacing3=6)
        txt.tag_config("img_label", foreground="#555555", font=("Arial", 10, "italic"),
                       spacing1=4, spacing3=8, justify="center")

        txt.insert("end", f"{categoria}\n", "tag_pill")
        txt.insert("end", f"{titulo}\n", "h1")
        txt.insert("end", f"{autor}  ·  {fecha}\n", "meta")
        if resumen:
            txt.insert("end", f"{resumen}\n", "resumen")
        txt.insert("end", "─" * 60 + "\n", "sep")

        self._renderizar_html(txt, cuerpo_html)
        txt.config(state="disabled")

    def _insertar_imagen_preview(self, txt, ruta_archivo):
        try:
            from PIL import Image as PILImage, ImageTk
            img = PILImage.open(ruta_archivo)
            max_w = 580
            if img.width > max_w:
                ratio = max_w / img.width
                img = img.resize((max_w, int(img.height * ratio)), PILImage.LANCZOS)
            photo = ImageTk.PhotoImage(img)
            self._imagenes_tk.append(photo)
            txt.image_create("end", image=photo, padx=4, pady=8)
            txt.insert("end", "\n")
        except Exception:
            txt.insert("end", "[📷 Imagen no disponible]\n", "img_label")

    def _renderizar_html(self, txt, html):
        import re as _re

        partes = _re.split(r'(<figure[^>]*>.*?</figure>)', html, flags=_re.DOTALL)

        for parte in partes:
            parte = parte.strip()
            if not parte:
                continue

            if parte.startswith('<figure'):
                src_match = _re.search(r'src="([^"]+)"', parte)
                if src_match:
                    src = src_match.group(1)
                    ruta_local = None
                    if self._carpeta_preview:
                        nombre    = Path(src).name
                        candidata = self._carpeta_preview / "img_inline" / nombre
                        if candidata.exists():
                            ruta_local = candidata
                    if ruta_local:
                        self._insertar_imagen_preview(txt, ruta_local)
                    else:
                        txt.insert("end", "[📷 Imagen del artículo]\n", "img_label")
                continue

            if parte.startswith('<div class="table-wrap">'):
                hdrs = _re.findall(r'<th[^>]*>(.*?)</th>', parte, _re.DOTALL)
                if hdrs:
                    hdr_text = " | ".join(_re.sub(r'<[^>]+>', '', h).strip() for h in hdrs)
                    txt.insert("end", hdr_text + "\n", "tabla_hdr")
                filas = _re.findall(r'<tr>(.*?)</tr>', parte, _re.DOTALL)
                for fila in filas:
                    celdas = _re.findall(r'<td[^>]*>(.*?)</td>', fila, _re.DOTALL)
                    if len(celdas) >= 2:
                        izq = _re.sub(r'<[^>]+>', '', celdas[0]).strip()
                        der = _re.sub(r'<[^>]+>', '', celdas[1]).strip()
                        if izq:
                            txt.insert("end", izq + "\n", "tabla_izq")
                        if der:
                            txt.insert("end", der + "\n", "tabla_der")
                continue

            bloques = parte.split("\n\n")
            for bloque in bloques:
                bloque = bloque.strip()
                if not bloque:
                    continue
                if bloque.startswith("<h2"):
                    texto = _re.sub(r'<[^>]+>', '', bloque)
                    txt.insert("end", texto + "\n", "h2")
                elif bloque.startswith("<h3"):
                    texto = _re.sub(r'<[^>]+>', '', bloque)
                    txt.insert("end", texto + "\n", "h3")
                else:
                    texto = _re.sub(r'<[^>]+>', '', bloque)
                    if texto:
                        txt.insert("end", texto + "\n", "p")

    def _confirmar(self):
        self.resultado = True
        self.grab_release()
        self.destroy()

    def _cancelar(self):
        self.resultado = False
        self.grab_release()
        self.destroy()


# ══════════════════════════════════════════════════════════════════════════════
#  VENTANA DE LOG
# ══════════════════════════════════════════════════════════════════════════════

class VentanaLog(ctk.CTkToplevel):
    def __init__(self, master, titulo="Construyendo el sitio..."):
        super().__init__(master)
        self.title(titulo)
        self.geometry("620x380")
        self.resizable(True, True)

        ctk.CTkLabel(self, text=titulo, font=("Arial", 13, "bold")).pack(pady=(14, 6), padx=16)

        self.barra = ctk.CTkProgressBar(self, mode="indeterminate", width=560)
        self.barra.pack(pady=(0, 10))
        self.barra.start()

        import tkinter as tk
        frame = ctk.CTkFrame(self, fg_color="#1a1a1a", corner_radius=8)
        frame.pack(fill="both", expand=True, padx=16, pady=(0, 14))

        self._txt = tk.Text(
            frame, bg="#1a1a1a", fg="#00ff88", font=("Courier", 11),
            relief="flat", padx=10, pady=10, state="disabled", wrap="word"
        )
        sb = tk.Scrollbar(frame, command=self._txt.yview)
        self._txt.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        self._txt.pack(fill="both", expand=True)

        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", lambda: None)

    def agregar(self, linea):
        self._txt.configure(state="normal")
        self._txt.insert("end", linea)
        self._txt.see("end")
        self._txt.configure(state="disabled")

    def finalizar(self, exito):
        self.barra.stop()
        self.barra.configure(mode="determinate")
        self.barra.set(1.0)
        self.barra.configure(progress_color="#0B7439" if exito else "#8E0209")
        self.protocol("WM_DELETE_WINDOW", self.destroy)
        ctk.CTkButton(self, text="Cerrar", command=self.destroy, height=34).pack(pady=(0, 10))


# ══════════════════════════════════════════════════════════════════════════════
#  BUILD
# ══════════════════════════════════════════════════════════════════════════════

def correr_build_con_log(ventana_log, callback_fin):
    def _hilo():
        proc = subprocess.Popen(
            ["python3", "build.py"], cwd=ROOT,
            stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
            text=True, bufsize=1,
        )
        for linea in proc.stdout:
            ventana_log.after(0, ventana_log.agregar, linea)
        proc.wait()
        exito = proc.returncode == 0
        ventana_log.after(0, ventana_log.finalizar, exito)
        ventana_log.after(0, callback_fin, exito, proc.returncode)

    threading.Thread(target=_hilo, daemon=True).start()


# ══════════════════════════════════════════════════════════════════════════════
#  APP PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Generador de Noticias — MJR-21")
        self.geometry("660x820")
        self.ultimo_slug = None

        ctk.CTkLabel(self, text="Generador de Noticias", font=("Arial", 20, "bold")).pack(pady=(20, 4))
        ctk.CTkLabel(self, text="MJR-21 — Departamento de Periodismo", text_color="gray").pack(pady=(0, 14))

        self.tabs = ctk.CTkTabview(self, width=600, height=700)
        self.tabs.pack(padx=20, pady=10, fill="both", expand=True)
        self.tabs.add("Crear noticia")
        self.tabs.add("Editar noticia")
        self.tabs.add("Eliminar / deshacer")

        self._construir_tab_crear(self.tabs.tab("Crear noticia"))
        self._construir_tab_editar(self.tabs.tab("Editar noticia"))
        self._construir_tab_eliminar(self.tabs.tab("Eliminar / deshacer"))

    def _campo(self, padre, placeholder, pady=5):
        entry = ctk.CTkEntry(padre, placeholder_text=placeholder, height=38)
        entry.pack(pady=pady, fill="x", padx=20)
        return entry

    # ── TAB CREAR ─────────────────────────────────────────────────────────

    def _construir_tab_crear(self, padre):
        self.c_ruta_doc    = None
        self.c_ruta_imagen = None
        self._carpeta_preview_temp = None

        self.c_btn_doc = ctk.CTkButton(
            padre, text="📄 Elegir archivo Word (.docx) o PDF",
            command=self._c_elegir_doc
        )
        self.c_btn_doc.pack(pady=(14, 4), fill="x", padx=20)
        self.c_lbl_doc = ctk.CTkLabel(padre, text="Ningún archivo elegido", text_color="gray")
        self.c_lbl_doc.pack()

        self.c_btn_imagen = ctk.CTkButton(
            padre, text="🖼️ Elegir foto de portada",
            command=self._c_elegir_imagen
        )
        self.c_btn_imagen.pack(pady=(8, 4), fill="x", padx=20)
        self.c_lbl_imagen = ctk.CTkLabel(padre, text="Ninguna foto elegida", text_color="gray")
        self.c_lbl_imagen.pack()

        self.c_entry_titulo    = self._campo(padre, "Título")
        self.c_entry_categoria = self._campo(padre, "Categoría (ej. Comunicados, Análisis...)")
        self.c_entry_autor     = self._campo(padre, "Autor")
        self.c_entry_fecha     = self._campo(padre, "Fecha (YYYY-MM-DD)")
        self.c_entry_fecha.insert(0, date.today().isoformat())
        self.c_entry_resumen   = self._campo(padre, "Resumen (1-2 líneas)")

        self.c_lbl_slug = ctk.CTkLabel(padre, text="", text_color="gray", font=("Arial", 11))
        self.c_lbl_slug.pack(pady=(4, 2))
        self.c_lbl_fecha_error = ctk.CTkLabel(padre, text="", text_color="#CB030E", font=("Arial", 11))
        self.c_lbl_fecha_error.pack()

        self.c_entry_titulo.bind("<KeyRelease>", self._c_actualizar_slug)
        self.c_entry_fecha.bind("<KeyRelease>",  self._c_validar_fecha)

        self.c_btn_previsualizar = ctk.CTkButton(
            padre, text="👁  Previsualizar antes de publicar",
            command=self._c_previsualizar,
            height=42, font=("Arial", 13, "bold"),
            fg_color="#0C4988", hover_color="#082F58"
        )
        self.c_btn_previsualizar.pack(pady=(14, 4), fill="x", padx=20)

        self.c_lbl_estado = ctk.CTkLabel(padre, text="", text_color="green", wraplength=540)
        self.c_lbl_estado.pack(pady=6)

    def _c_elegir_doc(self):
        ruta = elegir_archivo("Elegir noticia (Word o PDF)", FILTRO_DOC_ZENITY, FILTRO_DOC_TK)
        if not ruta:
            return
        self.c_ruta_doc = Path(ruta)
        self.c_lbl_doc.configure(text=f"{self.c_ruta_doc.name}  [{self.c_ruta_doc.suffix.upper()}]")
        titulo_sugerido = _obtener_titulo_sugerido(self.c_ruta_doc)
        if titulo_sugerido and not self.c_entry_titulo.get():
            self.c_entry_titulo.insert(0, titulo_sugerido)
        self._c_actualizar_slug()

    def _c_elegir_imagen(self):
        ruta = elegir_archivo("Elegir foto de portada", FILTRO_IMG_ZENITY, FILTRO_IMG_TK)
        if not ruta:
            return
        self.c_ruta_imagen = Path(ruta)
        self.c_lbl_imagen.configure(text=self.c_ruta_imagen.name)

    def _c_actualizar_slug(self, *_):
        fecha  = self.c_entry_fecha.get().strip() or date.today().isoformat()
        titulo = self.c_entry_titulo.get().strip()
        slug   = f"{fecha}-{slugify(titulo)}" if titulo else fecha
        self.c_lbl_slug.configure(text=f"Carpeta: nuevos/{slug}/")

    def _c_validar_fecha(self, *_):
        texto = self.c_entry_fecha.get().strip()
        if texto and not validar_fecha(texto):
            self.c_lbl_fecha_error.configure(text="⚠ Formato inválido. Usa YYYY-MM-DD")
        else:
            self.c_lbl_fecha_error.configure(text="")
        self._c_actualizar_slug()

    def _c_previsualizar(self):
        if not self.c_ruta_doc:
            messagebox.showerror("Falta el documento", "Elige primero un archivo .docx o .pdf.")
            return
        titulo    = self.c_entry_titulo.get().strip()
        categoria = self.c_entry_categoria.get().strip()
        autor     = self.c_entry_autor.get().strip()
        fecha     = self.c_entry_fecha.get().strip()
        resumen   = self.c_entry_resumen.get().strip()

        if not titulo or not categoria or not autor or not fecha:
            messagebox.showerror("Faltan datos", "Completa título, categoría, autor y fecha.")
            return
        if not validar_fecha(fecha):
            messagebox.showerror("Fecha inválida", "Usa el formato YYYY-MM-DD.")
            return

        import tempfile
        carpeta_temp = Path(tempfile.mkdtemp(prefix="mjr21_preview_"))
        self._carpeta_preview_temp = carpeta_temp

        slug_preview = slugify(titulo)
        cuerpo_html  = _convertir_a_cuerpo(self.c_ruta_doc, carpeta_temp, slug_preview)

        preview = VentanaPreview(
            self, titulo, categoria, autor, fecha, resumen, cuerpo_html,
            carpeta_preview=carpeta_temp
        )
        self.wait_window(preview)

        try:
            shutil.rmtree(carpeta_temp)
        except Exception:
            pass
        self._carpeta_preview_temp = None

        if preview.resultado:
            self._c_generar(titulo, categoria, autor, fecha, resumen)

    def _c_generar(self, titulo, categoria, autor, fecha, resumen):
        slug    = f"{fecha}-{slugify(titulo)}"
        carpeta = NUEVOS / slug

        if carpeta.exists():
            if not messagebox.askyesno("Ya existe", f"Ya existe '{slug}'. ¿Sobrescribir?"):
                return

        carpeta.mkdir(parents=True, exist_ok=True)
        cuerpo_html = _convertir_a_cuerpo(self.c_ruta_doc, carpeta, slug)

        contenido = (
            f"TITULO: {titulo}\n"
            f"CATEGORIA: {categoria}\n"
            f"AUTOR: {autor}\n"
            f"FECHA: {fecha}\n"
            f"RESUMEN: {resumen}\n"
            f"---\n"
            f"{cuerpo_html}\n"
        )
        (carpeta / "articulo.txt").write_text(contenido, encoding="utf-8")

        if self.c_ruta_imagen:
            shutil.copy(self.c_ruta_imagen, carpeta / f"foto{self.c_ruta_imagen.suffix.lower()}")

        self.ultimo_slug = slug
        self.c_btn_previsualizar.configure(state="disabled", text="⏳ Construyendo...")
        self.c_lbl_estado.configure(text="")

        log = VentanaLog(self, "Construyendo el sitio...")

        def al_terminar(exito, codigo):
            self.c_btn_previsualizar.configure(state="normal", text="👁  Previsualizar antes de publicar")
            if exito:
                self.c_lbl_estado.configure(
                    text=f"✅ '{slug}' publicada. Recarga el navegador.", text_color="green"
                )
                self._e_refrescar_lista()
                self._d_refrescar_lista()
            else:
                self.c_lbl_estado.configure(
                    text=f"⚠️ build.py terminó con error (código {codigo}).", text_color="red"
                )

        correr_build_con_log(log, al_terminar)

    # ── TAB EDITAR ────────────────────────────────────────────────────────

    def _construir_tab_editar(self, padre):
        self.e_slug_actual = None
        self.e_ruta_doc    = None
        self.e_ruta_imagen = None

        ctk.CTkLabel(
            padre, text="Elige una noticia para editar sus datos o reemplazar archivos.",
            text_color="gray", wraplength=520
        ).pack(pady=(14, 8))

        self.e_menu = ctk.CTkOptionMenu(padre, values=["(sin noticias)"], command=self._e_cargar)
        self.e_menu.pack(fill="x", padx=20, pady=(0, 4))
        ctk.CTkButton(padre, text="🔄 Actualizar lista", command=self._e_refrescar_lista, height=30).pack(pady=(0, 10))

        self.e_btn_doc = ctk.CTkButton(
            padre, text="📄 Reemplazar Word/PDF (opcional)", command=self._e_elegir_doc
        )
        self.e_btn_doc.pack(fill="x", padx=20, pady=(0, 4))
        self.e_lbl_doc = ctk.CTkLabel(
            padre, text="Sin cambio — se conserva el contenido actual",
            text_color="gray", font=("Arial", 11)
        )
        self.e_lbl_doc.pack()

        self.e_btn_imagen = ctk.CTkButton(
            padre, text="🖼️ Reemplazar foto de portada (opcional)", command=self._e_elegir_imagen
        )
        self.e_btn_imagen.pack(fill="x", padx=20, pady=(6, 4))
        self.e_lbl_imagen = ctk.CTkLabel(
            padre, text="Sin cambio — se conserva la foto actual",
            text_color="gray", font=("Arial", 11)
        )
        self.e_lbl_imagen.pack()

        self.e_entry_titulo    = self._campo(padre, "Título")
        self.e_entry_categoria = self._campo(padre, "Categoría")
        self.e_entry_autor     = self._campo(padre, "Autor")
        self.e_entry_fecha     = self._campo(padre, "Fecha (YYYY-MM-DD)")
        self.e_entry_resumen   = self._campo(padre, "Resumen")

        self.e_lbl_fecha_error = ctk.CTkLabel(padre, text="", text_color="#CB030E", font=("Arial", 11))
        self.e_lbl_fecha_error.pack()
        self.e_entry_fecha.bind("<KeyRelease>", self._e_validar_fecha)

        self.e_btn_guardar = ctk.CTkButton(
            padre, text="💾 Guardar cambios y reconstruir",
            command=self._e_guardar,
            height=42, font=("Arial", 13, "bold"),
            fg_color="#0B7439", hover_color="#075226"
        )
        self.e_btn_guardar.pack(pady=14, fill="x", padx=20)

        self.e_lbl_estado = ctk.CTkLabel(padre, text="", text_color="green", wraplength=540)
        self.e_lbl_estado.pack(pady=4)

        self._e_refrescar_lista()

    def _e_refrescar_lista(self):
        self._e_mapa = {f"{t}  ({s})": s for s, t in listar_articulos_existentes()}
        opciones = list(self._e_mapa.keys()) or ["(sin noticias)"]
        self.e_menu.configure(values=opciones)
        self.e_menu.set(opciones[0])
        if opciones[0] != "(sin noticias)":
            self._e_cargar(opciones[0])

    def _e_cargar(self, seleccion):
        slug = self._e_mapa.get(seleccion)
        if not slug:
            return
        self.e_slug_actual = slug
        self.e_ruta_doc    = None
        self.e_ruta_imagen = None
        self.e_lbl_doc.configure(text="Sin cambio — se conserva el contenido actual")
        self.e_lbl_imagen.configure(text="Sin cambio — se conserva la foto actual")

        meta = leer_metadata_articulo(slug)
        for entry, clave in [
            (self.e_entry_titulo,    "titulo"),
            (self.e_entry_categoria, "categoria"),
            (self.e_entry_autor,     "autor"),
            (self.e_entry_fecha,     "fecha"),
            (self.e_entry_resumen,   "resumen"),
        ]:
            entry.delete(0, "end")
            entry.insert(0, meta.get(clave, ""))

        self.e_lbl_estado.configure(text="")
        self.e_lbl_fecha_error.configure(text="")

    def _e_elegir_doc(self):
        ruta = elegir_archivo("Elegir nuevo Word o PDF", FILTRO_DOC_ZENITY, FILTRO_DOC_TK)
        if not ruta:
            return
        self.e_ruta_doc = Path(ruta)
        self.e_lbl_doc.configure(text=f"Nuevo archivo: {self.e_ruta_doc.name}  [{self.e_ruta_doc.suffix.upper()}]")

    def _e_elegir_imagen(self):
        ruta = elegir_archivo("Elegir nueva foto", FILTRO_IMG_ZENITY, FILTRO_IMG_TK)
        if not ruta:
            return
        self.e_ruta_imagen = Path(ruta)
        self.e_lbl_imagen.configure(text=f"Nueva foto: {self.e_ruta_imagen.name}")

    def _e_validar_fecha(self, *_):
        texto = self.e_entry_fecha.get().strip()
        if texto and not validar_fecha(texto):
            self.e_lbl_fecha_error.configure(text="⚠ Formato inválido. Usa YYYY-MM-DD")
        else:
            self.e_lbl_fecha_error.configure(text="")

    def _e_guardar(self):
        if not self.e_slug_actual:
            messagebox.showinfo("Nada seleccionado", "Elige primero una noticia de la lista.")
            return

        titulo    = self.e_entry_titulo.get().strip()
        categoria = self.e_entry_categoria.get().strip()
        autor     = self.e_entry_autor.get().strip()
        fecha     = self.e_entry_fecha.get().strip()
        resumen   = self.e_entry_resumen.get().strip()

        if not titulo or not categoria or not autor or not fecha:
            messagebox.showerror("Faltan datos", "Completa todos los campos.")
            return
        if not validar_fecha(fecha):
            messagebox.showerror("Fecha inválida", "Usa el formato YYYY-MM-DD.")
            return

        carpeta = NUEVOS / self.e_slug_actual

        if self.e_ruta_doc:
            inline_vieja = carpeta / "img_inline"
            if inline_vieja.exists():
                shutil.rmtree(inline_vieja)
            cuerpo_html = _convertir_a_cuerpo(self.e_ruta_doc, carpeta, self.e_slug_actual)
        else:
            txt_actual = (carpeta / "articulo.txt").read_text(encoding="utf-8")
            _, _, cuerpo_html = txt_actual.partition("---")
            cuerpo_html = cuerpo_html.strip()

        contenido = (
            f"TITULO: {titulo}\n"
            f"CATEGORIA: {categoria}\n"
            f"AUTOR: {autor}\n"
            f"FECHA: {fecha}\n"
            f"RESUMEN: {resumen}\n"
            f"---\n"
            f"{cuerpo_html}\n"
        )
        (carpeta / "articulo.txt").write_text(contenido, encoding="utf-8")

        if self.e_ruta_imagen:
            for vieja in carpeta.glob("foto.*"):
                vieja.unlink()
            shutil.copy(self.e_ruta_imagen, carpeta / f"foto{self.e_ruta_imagen.suffix.lower()}")

        self.e_btn_guardar.configure(state="disabled", text="⏳ Construyendo...")

        log = VentanaLog(self, "Aplicando cambios...")

        def al_terminar(exito, codigo):
            self.e_btn_guardar.configure(state="normal", text="💾 Guardar cambios y reconstruir")
            if exito:
                self.e_lbl_estado.configure(
                    text=f"✅ '{self.e_slug_actual}' actualizada. Recarga el navegador.",
                    text_color="green"
                )
                self._e_refrescar_lista()
                self._d_refrescar_lista()
            else:
                self.e_lbl_estado.configure(
                    text=f"⚠️ build.py falló (código {codigo}).", text_color="red"
                )

        correr_build_con_log(log, al_terminar)

    # ── TAB ELIMINAR ──────────────────────────────────────────────────────

    def _construir_tab_eliminar(self, padre):
        ctk.CTkLabel(
            padre, text="Elimina una noticia existente, o deshaz la última que creaste.",
            text_color="gray", wraplength=520
        ).pack(pady=(14, 10))

        self.d_lbl_ultimo = ctk.CTkLabel(
            padre, text="Última noticia creada en esta sesión: ninguna", wraplength=520
        )
        self.d_lbl_ultimo.pack(pady=(0, 6))

        self.d_btn_deshacer = ctk.CTkButton(
            padre, text="↩️ Deshacer la última noticia creada",
            command=self._d_deshacer, fg_color="#B98C02", hover_color="#8a6800"
        )
        self.d_btn_deshacer.pack(pady=10, fill="x", padx=20)

        ctk.CTkLabel(padre, text="— o elige cualquier noticia para eliminarla —",
                     text_color="gray").pack(pady=(16, 6))

        self.d_menu = ctk.CTkOptionMenu(padre, values=["(sin noticias)"])
        self.d_menu.pack(pady=6, fill="x", padx=20)
        ctk.CTkButton(padre, text="🔄 Actualizar lista", command=self._d_refrescar_lista, height=30).pack(pady=4)

        self.d_btn_eliminar = ctk.CTkButton(
            padre, text="🗑️ Eliminar la noticia seleccionada",
            command=self._d_eliminar, fg_color="#8E0209", hover_color="#5e0106"
        )
        self.d_btn_eliminar.pack(pady=14, fill="x", padx=20)

        self.d_lbl_estado = ctk.CTkLabel(padre, text="", text_color="green", wraplength=520)
        self.d_lbl_estado.pack(pady=6)

        self._d_refrescar_lista()

    def _d_refrescar_lista(self):
        self._d_mapa = {f"{t}  ({s})": s for s, t in listar_articulos_existentes()}
        opciones = list(self._d_mapa.keys()) or ["(sin noticias)"]
        self.d_menu.configure(values=opciones)
        self.d_menu.set(opciones[0])
        if self.ultimo_slug:
            self.d_lbl_ultimo.configure(text=f"Última noticia creada: {self.ultimo_slug}")
        else:
            self.d_lbl_ultimo.configure(text="Última noticia creada en esta sesión: ninguna")

    def _d_deshacer(self):
        if not self.ultimo_slug:
            messagebox.showinfo("Nada que deshacer", "No has creado ninguna noticia en esta sesión.")
            return
        if not messagebox.askyesno("Confirmar", f"¿Eliminar '{self.ultimo_slug}'?"):
            return
        eliminar_articulo(self.ultimo_slug)
        slug = self.ultimo_slug
        self.ultimo_slug = None
        self._d_ejecutar_build_eliminar(slug)

    def _d_eliminar(self):
        seleccion = self.d_menu.get()
        slug = self._d_mapa.get(seleccion)
        if not slug:
            messagebox.showinfo("Nada seleccionado", "No hay ninguna noticia para eliminar.")
            return
        if not messagebox.askyesno(
            "Confirmar eliminación",
            f"¿Eliminar permanentemente '{seleccion}'?\n\nEsto no se puede deshacer."
        ):
            return
        eliminar_articulo(slug)
        if self.ultimo_slug == slug:
            self.ultimo_slug = None
        self._d_ejecutar_build_eliminar(slug)

    def _d_ejecutar_build_eliminar(self, slug):
        self.d_btn_eliminar.configure(state="disabled")
        log = VentanaLog(self, "Reconstruyendo tras eliminación...")

        def al_terminar(exito, codigo):
            self.d_btn_eliminar.configure(state="normal")
            if exito:
                self.d_lbl_estado.configure(
                    text=f"✅ '{slug}' eliminada y sitio reconstruido. Recarga el navegador.",
                    text_color="green"
                )
            else:
                self.d_lbl_estado.configure(
                    text=f"⚠️ Se borró la carpeta pero build.py falló (código {codigo}).",
                    text_color="red"
                )
            self._d_refrescar_lista()
            self._e_refrescar_lista()

        correr_build_con_log(log, al_terminar)


# ══════════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    NUEVOS.mkdir(parents=True, exist_ok=True)
    app = App()
    app.mainloop()